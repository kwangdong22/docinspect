from fastapi import FastAPI, UploadFile, File
import io
import docx
import fitz  # PyMuPDF
import re

app = FastAPI()

@app.get("/health")
def health_check():
    return {"status": "extractor_ok"}

@app.post("/extract")
async def extract_document(file: UploadFile = File(...)):
    content = await file.read()
    text = ""
    page_count = 0
    has_tables = False
    has_images = False
    
    if file.filename.endswith(".pdf"):
        pdf_doc = fitz.open(stream=content, filetype="pdf")
        page_count = len(pdf_doc)
        for page in pdf_doc:
            text += page.get_text()
            # Check for images on the page
            if len(page.get_images()) > 0:
                has_images = True
            
    elif file.filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(content))
        page_count = 1 
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Check for tables and images in Word
        if len(doc.tables) > 0:
            has_tables = True
        if len(doc.inline_shapes) > 0:
            has_images = True

    words = text.split()
    word_count = len(words)
    character_count = len(text)

    # Regular Expressions to find PII
    # Finds standard email formats
    emails = list(set(re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)))
    
    # Finds phone numbers like (123) 456-7890 or 123-456-7890
    phones = list(set(re.findall(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)))
    
    # Finds http/https links
    links = list(set(re.findall(r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+', text)))

    return {
        "page_count": page_count,
        "word_count": word_count,
        "character_count": character_count,
        "emails": emails,
        "phones": phones,
        "links": links,
        "has_tables": has_tables,
        "has_images": has_images
    }
